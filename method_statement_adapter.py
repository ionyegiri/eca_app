# method_statement_adapter.py
# Adapter that parses project Method Statements (DOCX/PDF) and extracts default rules
# Author: Ikechukwu Onyegiri + Copilot

from __future__ import annotations
import re
from dataclasses import dataclass, field
from typing import List, Optional, Tuple, Dict, Union

# Third-party parsers (available in the sandbox runtime)
import fitz  # PyMuPDF
from docx import Document  # python-docx

from models import (
    DeterministicInputs, FADOption, JRCurve, JRCurvePoint, FCGRParams
)

# ----------------------------
# Defaults container
# ----------------------------

@dataclass
class MethodDefaults:
    # Global ECA policy
    fad_option: FADOption = FADOption.OPTION_2
    tearing_limit_fraction_t: float = 0.10              # DNV: accumulated tearing ≤ 10% WT
    tearing_increment_cap_frac: float = 2.0 / 3.0       # ≤ 2/3 of SENT max-load tearing per increment (DNV)
    rechar_rule_ligament_over_a: float = 0.5            # embedded → surface when ligament < 0.5·a

    # Reeling (ductile tearing)
    use_sent_jr_for_reeling: bool = True
    jr_curve_points: List[JRCurvePoint] = field(default_factory=list)  # optional: parsed from tables if present

    # FCGR policies (textual selection; numerical C,m can be set elsewhere)
    # e.g., 'BS7910_in_air_mean_plus_2std', 'project_env_with_KDF'
    fcgr_install_policy: str = "BS7910_in_air_mean_plus_2std"
    fcgr_operation_policy: str = "BS7910_in_air_mean_plus_2std"
    fcgr_operation_id_env_policy: Optional[str] = None  # e.g., 'aggressive_env_with_company_KDF'

    # Optional numerical overrides for Paris parameters (rarely present directly in MS)
    fcgr_install_params: Optional[FCGRParams] = None
    fcgr_operation_params: Optional[FCGRParams] = None

    notes: List[str] = field(default_factory=list)


# ----------------------------
# Text extraction
# ----------------------------

def _extract_text_from_docx(path: str) -> str:
    doc = Document(path)
    texts = []
    for p in doc.paragraphs:
        if p.text:
            texts.append(p.text)
    # Also include table text
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                val = cell.text.strip()
                if val:
                    texts.append(val)
    return "\n".join(texts)

def _extract_text_from_pdf(path: str) -> str:
    out = []
    with fitz.open(path) as doc:
        for page in doc:
            out.append(page.get_text("text"))
    return "\n".join(out)


# ----------------------------
# Heuristic parsers
# ----------------------------

_re_pct = re.compile(r"(\d+(?:\.\d+)?)\s*%")
_re_frac = re.compile(r"(\d+(?:\.\d+)?)/(\d+(?:\.\d+)?)")

def _find_tearing_limits(text: str, md: MethodDefaults) -> None:
    # Accumulated tearing ≤ 10% WT
    if re.search(r"(accumulated|total)\s+tearing.*(10\s*%|0\.1\s*t|0\.10\s*t)", text, re.IGNORECASE):
        md.tearing_limit_fraction_t = 0.10
        md.notes.append("Detected tearing limit ≤ 10% WT.")
    # Increment cap ~ 2/3 max-load tearing
    if re.search(r"(one\s+strain\s+increment|per\s+increment).*two[-\s]*thirds|2/3", text, re.IGNORECASE):
        md.tearing_increment_cap_frac = 2.0 / 3.0
        md.notes.append("Detected per-increment tearing cap ≤ 2/3 of SENT max-load tearing.")

def _find_fad_option(text: str, md: MethodDefaults) -> None:
    if re.search(r"Option\s*2.*(stress[-\s]*strain|material[-\s]*specific)\s*FAD", text, re.IGNORECASE):
        md.fad_option = FADOption.OPTION_2
        md.notes.append("Detected FAD Option 2 (material-specific from stress–strain).")
    elif re.search(r"Option\s*1", text, re.IGNORECASE):
        # In case some MS opts for Option 1 for certain checks
        md.fad_option = FADOption.OPTION_1
        md.notes.append("Detected FAD Option 1; verify project intent.")

def _find_fcgr_policies(text: str, md: MethodDefaults) -> None:
    # Installation: BS 7910 in-air (Mean+2σ)
    if re.search(r"(installation).*(BS\s*7910).*(in[-\s]*air).*(mean\s*\+\s*2\s*std)", text, re.IGNORECASE | re.DOTALL):
        md.fcgr_install_policy = "BS7910_in_air_mean_plus_2std"
        md.notes.append("Installation FCGR: BS 7910 In‑air (Mean+2σ).")
    # Operation: similar baseline; ID env / KDF possibility
    if re.search(r"(operation|service).*(BS\s*7910).*(in[-\s]*air).*(mean\s*\+\s*2\s*std)", text, re.IGNORECASE | re.DOTALL):
        md.fcgr_operation_policy = "BS7910_in_air_mean_plus_2std"
        md.notes.append("Operation FCGR: BS 7910 In‑air (Mean+2σ) detected.")
    if re.search(r"(ID|internal\s+diameter).*(aggressive|sour|environment).*KDF|acceleration\s+factor", text, re.IGNORECASE):
        md.fcgr_operation_id_env_policy = "aggressive_env_with_company_KDF"
        md.notes.append("ID environment with COMPANY KDF detected for operation phase.")

def _find_rechar_rule(text: str, md: MethodDefaults) -> None:
    # Embedded → surface re-characterization rule (ligament threshold)
    # Common rule: ligament < 0.5 * a
    if re.search(r"(embedded).*(re-?characteri[sz]e|recharacteri[sz]ation).*ligament.*0\.5\s*\*?\s*a", text, re.IGNORECASE | re.DOTALL):
        md.rechar_rule_ligament_over_a = 0.5
        md.notes.append("Embedded→Surface rule detected: ligament < 0.5·a.")

def _maybe_parse_jr_points_from_tables_docx(path: str) -> List[JRCurvePoint]:
    """
    Best-effort JR-curve table parsing from DOCX: look for columns like Δa, J.
    Returns list of JRCurvePoint.
    """
    pts: List[JRCurvePoint] = []
    try:
        doc = Document(path)
    except Exception:
        return pts
    for t in doc.tables:
        # Try to detect headers
        headers = [c.text.lower().strip() for c in t.rows[0].cells]
        has_da = any("da" in h or "Δa" in h or "delta a" in h for h in headers)
        has_j = any(h.startswith("j") for h in headers)
        if not (has_da and has_j):
            continue
        # Find column indices
        try:
            i_da = next(i for i, h in enumerate(headers) if ("da" in h or "Δa" in h or "delta a" in h))
            i_j = next(i for i, h in enumerate(headers) if (h.startswith("j")))
        except StopIteration:
            continue
        for r in t.rows[1:]:
            try:
                s_da = r.cells[i_da].text.strip().replace(",", "")
                s_j = r.cells[i_j].text.strip().replace(",", "")
                if not s_da or not s_j:
                    continue
                da = float(re.findall(r"[-+]?\d*\.?\d+", s_da)[0])
                Jv = float(re.findall(r"[-+]?\d*\.?\d+", s_j)[0])
                pts.append(JRCurvePoint(da=da, J=Jv))
            except Exception:
                continue
    return pts


# ----------------------------
# Public API
# ----------------------------

def parse_method_statement(path: str) -> MethodDefaults:
    """
    Parse a project Method Statement (DOCX or PDF) and return defaults.
    """
    text = ""
    path_lower = path.lower()
    if path_lower.endswith(".docx"):
        text = _extract_text_from_docx(path)
    elif path_lower.endswith(".pdf"):
        text = _extract_text_from_pdf(path)
    else:
        raise ValueError("Unsupported file type. Please provide a DOCX or PDF.")

    md = MethodDefaults()

    # Heuristic detections
    _find_tearing_limits(text, md)
    _find_fad_option(text, md)
    _find_fcgr_policies(text, md)
    _find_rechar_rule(text, md)

    # Optionally parse JR points from DOCX tables if present
    if path_lower.endswith(".docx"):
        jr_pts = _maybe_parse_jr_points_from_tables_docx(path)
        if jr_pts:
            md.jr_curve_points = jr_pts
            md.notes.append(f"Parsed {len(jr_pts)} JR-curve points from MS table.")

    return md


def apply_defaults_to_inputs(inputs: DeterministicInputs, md: MethodDefaults) -> DeterministicInputs:
    """
    Apply MethodDefaults to a DeterministicInputs object.
    Returned object is modified in place and also returned for convenience.
    """
    inputs.fad_option = md.fad_option
    inputs.tearing_limit_fraction_t = md.tearing_limit_fraction_t
    inputs.rechar_rule_ligament_over_a = md.rechar_rule_ligament_over_a
    # Reeling JR-curve policy
    if md.use_sent_jr_for_reeling:
        if md.jr_curve_points:
            inputs.jr_curve_reeling = JRCurve(points=md.jr_curve_points)
        else:
            # Keep existing JR if provided elsewhere (upload), otherwise None
            pass

    # FCGR policies: here we only set policy notes; numerical FCGR (C,m) can be set via UI or BS7910 presets
    if inputs.fcgr_install:
        inputs.fcgr_install.env_factor = 1.0  # In-air default; can be overridden
    if inputs.fcgr_operation:
        if md.fcgr_operation_id_env_policy == "aggressive_env_with_company_KDF":
            inputs.fcgr_operation.env_factor = max(1.0, inputs.fcgr_operation.env_factor)
        else:
            inputs.fcgr_operation.env_factor = 1.0

    return inputs


def summarize_defaults(md: MethodDefaults) -> Dict[str, Union[str, float, int]]:
    """
    Provide a compact summary suitable for the UI.
    """
    return {
        "FAD option": md.fad_option.value,
        "Tearing limit (fraction of WT)": md.tearing_limit_fraction_t,
        "Per-increment tearing cap (fraction of SENT max-load tearing)": md.tearing_increment_cap_frac,
        "Recharacterization rule (ligament < X * a)": md.rechar_rule_ligament_over_a,
        "Use SENT JR for reeling": md.use_sent_jr_for_reeling,
        "JR points parsed": len(md.jr_curve_points),
        "Installation FCGR policy": md.fcgr_install_policy,
        "Operation FCGR policy": md.fcgr_operation_policy,
        "Operation ID env policy": md.fcgr_operation_id_env_policy or "N/A",
        "Notes": " | ".join(md.notes) if md.notes else "—"
    }
